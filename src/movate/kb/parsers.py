"""Document parsers — uploaded file bytes → extracted text.

The KB ingest pipeline operates on plain text strings. Real KB
content arrives as PDFs (policy documents, runbooks), DOCXes
(internal wikis exported to Word), HTML pages (Confluence /
Notion exports). Without a parser layer, every operator has to
pre-extract text on their own before ``mdk kb ingest`` accepts
the content — a friction point that the PR-D Chainlit upload
flow made acutely visible (PDFs get ``status="skipped"``).

This module is the dispatch layer. Each format gets a small
parser; the orchestrator picks one by file extension. Failures
return ``None`` so the caller can surface a clean
``status="skipped"`` rather than 500-ing the upload.

v0.9 MVP scope:

* **PDF** via pypdf (PR-G) — handles text-based PDFs. When
  pypdf extracts no text (scanned-image PDFs), falls back to
  OCR via the optional ``[ocr]`` extra (pdf2image + pytesseract
  + system Tesseract binary). Install with
  ``pip install movate-cli[ocr]``. (PR-CC)
* **DOCX** via python-docx (PR-L) — handles standard Word
  documents. Honors heading levels as paragraph boundaries.
* **HTML** via readability-lxml + BeautifulSoup (PR-M) — runs
  Mozilla's Readability algorithm to extract just the main-article
  content, then strips tags. Right tier for Confluence / Notion /
  blog exports which carry heavy nav / sidebar / ad chrome.
* **Markdown / plain text** via UTF-8 decode (already supported
  by the runtime endpoint; included here for unified dispatch).

Out of scope (tracked in BACKLOG §10.6):

* **Legacy .doc** (binary Word format pre-2007) — python-docx
  rejects these; operators should convert to .docx first.
* **EPUB / PPTX / scanned-image content other than PDFs** —
  each needs its own parser tier; tackle by demand.

The :class:`DocumentParser` Protocol lays the groundwork for
per-format entry-point registration when more formats land — but
v0.9 keeps it simple with a dispatch dict.
"""

from __future__ import annotations

import io
import logging
import os
import re
from typing import NamedTuple, Protocol

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# OCR tunables — can be overridden at process level for non-default setups.
# ---------------------------------------------------------------------------

# DPI for PDF→image rasterisation. 300 is the ISO-recommended minimum for
# OCR; 200 DPI (the prior default) is too coarse for small fonts and
# compressed scans. Higher DPI (400-600) improves accuracy on very small
# text at the cost of memory + latency — 300 is the sweet spot for most
# KB ingest workloads.
_OCR_DPI: int = 300

# Tesseract engine + page-segmentation mode:
#   --oem 1  → LSTM neural engine (default for Tesseract 4+; more accurate
#              than legacy OCR engine)
#   --psm 6  → Assume a single uniform block of text (appropriate for scanned
#              pages that are mostly prose; use --psm 11 for sparse layouts)
_OCR_CONFIG: str = "--oem 1 --psm 6"


class ParseResult(NamedTuple):
    """Return value of :func:`parse_document` and each per-format parser.

    Carrying ``ocr_used`` alongside the text lets the ingest pipeline
    tag :class:`~movate.core.models.KbChunk` records with whether they
    were extracted from image content via Tesseract OCR rather than
    native text extraction. Operators can then filter / audit OCR-sourced
    chunks separately (e.g. to flag lower-confidence content for review).
    """

    text: str
    """Extracted text, stripped of leading/trailing whitespace."""

    ocr_used: bool
    """True iff the text was produced by Tesseract OCR rather than a
    native text-extraction path (pypdf text layer, docx paragraphs,
    readability HTML strip). Always False for non-PDF formats."""


# File extensions the dispatch layer accepts. Sorted by frequency
# of operator-uploaded format (rough — refine when we have telemetry).
# Anything not in this set returns ``None`` from
# :func:`parse_document` so the caller reports ``status="skipped"``.
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset(
    {
        ".md",
        ".markdown",
        ".txt",
        ".pdf",
        ".docx",
        ".html",
        ".htm",
    }
)


class DocumentParser(Protocol):
    """Per-format parser interface. Each registered parser converts
    raw bytes to a :class:`ParseResult` or returns ``None`` on failure.

    Implementations MUST NOT raise; they log and return ``None`` so
    the dispatch layer can route a single bad file's
    ``status="skipped"`` without sinking a multi-file upload.
    """

    def __call__(self, content: bytes) -> ParseResult | None: ...


def parse_document(filename: str, content: bytes) -> ParseResult | None:
    """Extract text from ``content`` based on ``filename``'s extension.

    Returns a :class:`ParseResult` on success, or ``None`` if:

    * the extension isn't in :data:`SUPPORTED_EXTENSIONS`, OR
    * the parser failed (corrupt bytes, encrypted PDF, etc).

    The two failure modes are distinguishable by the caller via
    :func:`is_supported_extension` — useful for status reporting
    (``"skipped"`` for unsupported vs ``"empty"`` for parsed-but-no-text).

    Args:
        filename: The uploaded file's name (used only for extension
            dispatch — content is identified purely by its bytes).
        content: The raw bytes from the upload.

    Returns:
        :class:`ParseResult` on success, ``None`` on failure.
    """
    if not filename:
        return None
    ext = _extension(filename)
    parser = _PARSERS.get(ext)
    if parser is None:
        return None
    return parser(content)


def is_supported_extension(filename: str) -> bool:
    """Return True if ``filename``'s extension has a registered parser.

    Used by the KB upload endpoint + CLI to short-circuit
    obviously-unsupported files before reading the bytes off the
    wire (saves bandwidth on bulk uploads).
    """
    return _extension(filename) in SUPPORTED_EXTENSIONS


def _extension(filename: str) -> str:
    """Return the lowercase extension including the dot, or ``""``."""
    idx = filename.rfind(".")
    if idx < 0:
        return ""
    return filename[idx:].lower()


# ---------------------------------------------------------------------------
# Per-format parsers
# ---------------------------------------------------------------------------


def parse_text(content: bytes) -> ParseResult | None:
    """Decode UTF-8 text. Returns ``None`` on decode failure.

    Used for .md / .markdown / .txt. Same semantics as the KB
    upload endpoint's existing inline decode — extracting here
    centralizes the failure-mode reporting.
    """
    try:
        return ParseResult(text=content.decode("utf-8"), ocr_used=False)
    except UnicodeDecodeError:
        return None


def parse_pdf(content: bytes) -> ParseResult | None:
    """Extract text from a PDF via pypdf, with OCR fallback for scanned images.

    Returns the concatenated page text (one page per ``\\n\\n``
    paragraph) so the downstream paragraph-chunker can split
    naturally on page boundaries. Returns ``None`` if:

    * pypdf rejects the bytes (corrupt / not a PDF / encrypted
      without a password), OR
    * every page yields empty text AND the ``[ocr]`` extra is not
      installed (scanned-image PDF — install
      ``pip install movate-cli[ocr]`` to enable Tesseract OCR), OR
    * OCR is available but also produces no text.

    When the optional ``[ocr]`` extra IS installed (pdf2image +
    pytesseract + system Tesseract binary), scanned-image pages are
    handled transparently: for each page where pypdf returns no text,
    a per-page Poppler rasterisation step (``_OCR_DPI`` DPI, default
    300) is followed by Tesseract OCR. This enables **mixed PDFs**
    where some pages are text-based and others are scanned images —
    both contribute to the extracted corpus. Callers see no difference
    — they just get text back regardless of PDF content type.

    Why pypdf vs alternatives:

    * **pypdf** (~250KB, pure Python): text extraction only,
      enough for the 80% case. Handles text-based PDFs reliably.
    * **pdfplumber** (~600KB + multiple transitive deps): adds
      table extraction. Overkill for the v0.9 MVP — operators
      who need table-aware extraction can preprocess.
    * **pdfminer.six** (~1MB): more accurate layout extraction
      but 3x larger install + slower per-page parse.

    For text-extraction-only on internal docs, pypdf is the
    correct tier.
    """
    try:
        import pypdf  # noqa: PLC0415 — lazy: only paid for when an operator uploads a PDF
    except ImportError:
        logger.warning(
            "pypdf is not installed — PDF parsing unavailable. Install via: uv pip install pypdf"
        )
        return None

    try:
        reader = pypdf.PdfReader(io.BytesIO(content))
    except Exception as exc:
        logger.warning("pypdf failed to open PDF: %s", exc)
        return None

    # Encrypted PDFs need a password to extract. We don't have a
    # password channel in the upload UI; surface as ``None`` so the
    # operator sees ``status="skipped"`` rather than a partial extract.
    if reader.is_encrypted:
        logger.warning("PDF is encrypted; skipping")
        return None

    pages: list[str] = []
    ocr_used = False
    for page_num, page in enumerate(reader.pages):
        try:
            text = page.extract_text()
        except Exception as exc:
            logger.warning("page %d extraction failed: %s; skipping page", page_num, exc)
            continue
        if text and text.strip():
            pages.append(text.strip())
        else:
            # Empty page — scanned image. Attempt per-page OCR so mixed
            # PDFs (text pages interleaved with scanned pages) are handled
            # correctly instead of silently dropping the scanned pages.
            ocr_text = _ocr_pdf(content, page_num=page_num)
            if ocr_text:
                pages.append(ocr_text)
                ocr_used = True

    if not pages:
        return None

    # Join pages with the paragraph boundary the chunker uses so
    # natural page breaks become chunk breaks.
    return ParseResult(text="\n\n".join(pages), ocr_used=ocr_used)


def _ocr_pdf(content: bytes, *, page_num: int) -> str | None:
    """OCR a single PDF page (0-based index) using pdf2image + pytesseract.

    Called by :func:`parse_pdf` for every page that yields no text from
    pypdf — this per-page design enables **mixed PDFs** (text and scanned
    pages interleaved) rather than only handling fully-scanned documents.

    Pipeline:

    1. **pdf2image** rasterises the requested page at ``_OCR_DPI`` DPI
       (300 by default) using Poppler's pdftoppm renderer. Passing
       ``first_page``/``last_page`` means only the target page is decoded,
       which avoids re-rasterising the entire PDF for every empty page.
    2. **pytesseract** runs Tesseract with ``_OCR_CONFIG``
       (``--oem 1 --psm 6`` by default): LSTM engine + single-block
       page-segmentation mode, which works well for dense prose pages.
    3. Whitespace is normalised: runs of spaces/tabs are collapsed to a
       single space, and three or more consecutive newlines are reduced
       to two, preserving paragraph structure without noisy blank lines.

    Language is read from the ``MOVATE_OCR_LANG`` environment variable
    (default ``"eng"``). Set it to a Tesseract language code (e.g.
    ``"fra"`` for French, ``"deu"`` for German, ``"eng+fra"`` for
    multilingual) before starting the ingest process.

    Returns the normalised page text, or ``None`` if:

    * ``pdf2image`` or ``pytesseract`` are not installed (``[ocr]``
      extra missing) — logged at DEBUG so the import absence stays quiet.
    * Poppler is not installed (pdf2image raises) — logged at WARNING.
    * Tesseract is not installed or fails — logged at WARNING.
    * OCR succeeds but produces only whitespace.

    Intentionally synchronous — ``parse_pdf`` is also sync; large-batch
    ingest callers off-load the whole parse call to a thread/executor
    at a higher level.
    """
    try:
        import pdf2image  # type: ignore[import-not-found]  # noqa: PLC0415
        import pytesseract  # type: ignore[import-not-found]  # noqa: PLC0415
    except ImportError:
        logger.debug(
            "OCR deps not installed (pdf2image / pytesseract). "
            "Install movate-cli[ocr] to enable OCR for scanned PDFs."
        )
        return None

    try:
        images = pdf2image.convert_from_bytes(
            content,
            dpi=_OCR_DPI,
            first_page=page_num + 1,
            last_page=page_num + 1,
        )
    except Exception as exc:
        logger.warning("pdf2image failed to rasterise page %d for OCR: %s", page_num, exc)
        return None

    if not images:
        return None

    lang = os.environ.get("MOVATE_OCR_LANG", "eng")
    try:
        text: str = pytesseract.image_to_string(images[0], lang=lang, config=_OCR_CONFIG)
    except Exception as exc:
        logger.warning("tesseract OCR failed on page %d: %s", page_num, exc)
        return None

    # Normalise whitespace: collapse tab/space runs; trim excess blank lines.
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip() or None


def parse_docx(content: bytes) -> ParseResult | None:
    """Extract text from a .docx via python-docx.

    Returns paragraphs joined by ``\\n\\n`` so the downstream chunker
    treats them as natural chunk boundaries. Inserts a paragraph
    break around any paragraph with the "Heading 1" / "Heading 2" /
    etc. style so heading-style cues survive the round-trip even if
    the heading itself doesn't get pulled into its own chunk.

    Returns ``None`` if:

    * python-docx rejects the bytes (not a valid .docx — legacy
      binary .doc, corrupt zip, or a misnamed file), OR
    * the document has no extractable text (paragraphs all empty,
      or document contains only embedded images / tables we don't
      extract from in v0.9).

    Why python-docx vs alternatives:

    * **python-docx** (~3MB incl. lxml): mature, well-maintained,
      handles standard .docx files. Only reads from the document's
      main body; doesn't touch headers/footers (intentional — they
      are usually boilerplate noise).
    * **mammoth**: better at converting to HTML/Markdown but ~6MB
      with extra deps. Overkill for plain text extraction.
    * **textract**: heavyweight (depends on many system tools).

    Tables in .docx files are NOT extracted in v0.9 — most real KB
    content lives in paragraphs, and table cells would need their
    own row/col reconstruction logic. Operator can convert tables
    to inline text in Word before upload if they're critical.

    ``ocr_used`` is always ``False`` — DOCX text extraction is
    native (no OCR path exists for this format in v0.9).
    """
    try:
        import docx  # noqa: PLC0415 — lazy: only paid for when an operator uploads a .docx
    except ImportError:
        logger.warning(
            "python-docx is not installed — DOCX parsing unavailable. "
            "Install via: uv pip install python-docx"
        )
        return None

    try:
        # python-docx's Document() accepts a file-like object.
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:
        logger.warning("python-docx failed to open document: %s", exc)
        return None

    paragraphs: list[str] = []
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        # Headings get their own paragraph so the chunker doesn't
        # blur a "Refund Policy" heading into the surrounding
        # body text — operator-visible structure stays intact.
        style_name = getattr(getattr(para, "style", None), "name", "") or ""
        if style_name.lower().startswith("heading"):
            paragraphs.append(text)
        else:
            paragraphs.append(text)

    if not paragraphs:
        # Empty document OR contained only tables / images we don't
        # extract from. Either way the result is unusable.
        return None

    # Same join semantics as parse_pdf — chunker splits on \n\n.
    return ParseResult(text="\n\n".join(paragraphs), ocr_used=False)


def parse_html(content: bytes) -> ParseResult | None:
    """Extract main-article text from HTML via Readability + BeautifulSoup.

    Two-stage pipeline:

    1. **Readability** — runs Mozilla's algorithm against the DOM,
       extracts the ``<article>`` / main-content block, throws away
       nav / sidebar / footer / ads. Same heuristic browsers' Reader
       View uses. This step is the difference between "useful for
       retrieval" and "garbage with menu items embedded in every
       chunk".
    2. **BeautifulSoup strip-tags** on the extracted block, with
       paragraph-break preservation between block-level elements
       (``<p>``, ``<h1>``-``<h6>``, ``<li>``, ``<br>``).

    Returns ``None`` if:

    * The bytes don't decode as text (rare — HTML is text-based, but
      arbitrary binary uploads with a ``.html`` extension hit this).
    * Readability can't find a main-content block AND the strip-tags
      fallback produces empty output (e.g. ``<html></html>``).

    Why readability + BeautifulSoup vs alternatives:

    * **Naive strip-all-tags**: keeps the whole page including nav /
      sidebar / footer junk. Useless for retrieval over real
      Confluence / Notion exports.
    * **trafilatura** (~600KB + deps): broader feature set
      (boilerplate detection, metadata extraction). Right tier for
      a content-pipeline product, overkill for KB ingest.
    * **html2text**: produces clean Markdown but ~100KB + slower
      to extract main content. We don't need Markdown round-trip —
      plain text feeds the chunker fine.

    ``ocr_used`` is always ``False`` — HTML text extraction is
    native (no OCR path for HTML in v0.9).
    """
    try:
        from bs4 import BeautifulSoup  # noqa: PLC0415 — lazy: only paid for HTML uploads
        from readability import Document  # noqa: PLC0415
    except ImportError:
        logger.warning(
            "readability-lxml or beautifulsoup4 not installed — HTML parsing unavailable. "
            "Install via: uv pip install readability-lxml beautifulsoup4"
        )
        return None

    # Decode bytes. HTML in the wild uses many encodings (latin-1,
    # cp1252, gbk, ...); try utf-8 first, fall back to latin-1 (which
    # never fails). The lossy fallback is acceptable because we're
    # extracting for retrieval, not round-tripping the original.
    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = content.decode("latin-1")
        except UnicodeDecodeError:
            return None

    # Stage 1: Readability picks the main-article block.
    main_html: str
    try:
        doc = Document(text)
        main_html = doc.summary(html_partial=True)
    except Exception as exc:
        logger.warning("readability failed: %s; falling back to raw strip", exc)
        main_html = text

    # Stage 2: BeautifulSoup strip-tags with block-element paragraph
    # breaks. Replace block-level tags with `\n\n` before text
    # extraction so paragraphs / headings / list items become natural
    # chunk boundaries.
    try:
        soup = BeautifulSoup(main_html, "html.parser")
    except Exception as exc:
        logger.warning("BeautifulSoup failed to parse HTML: %s", exc)
        return None

    # Drop script / style / noscript outright — even when readability
    # selects them, they're never useful retrieval content.
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    # Insert paragraph breaks before block-level elements so the
    # chunker's \n\n splitter respects HTML's visual structure.
    for tag_name in ("p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "br", "div"):
        for tag in soup.find_all(tag_name):
            tag.insert_before("\n\n")

    raw = soup.get_text()
    # Collapse runs of whitespace within paragraphs; preserve the
    # \n\n paragraph boundaries between them.
    paragraphs: list[str] = []
    for para in raw.split("\n\n"):
        # Normalize internal whitespace to single spaces — preserves
        # readability without keeping HTML's stray tab/newline noise.
        cleaned = " ".join(para.split())
        if cleaned:
            paragraphs.append(cleaned)

    if not paragraphs:
        return None
    return ParseResult(text="\n\n".join(paragraphs), ocr_used=False)


# Dispatch table — extension → parser. Keeping this at the module
# bottom lets each parser's docstring sit next to the function for
# readability. Extending: add the extension to SUPPORTED_EXTENSIONS,
# add the parser function above, register here.
_PARSERS: dict[str, DocumentParser] = {
    ".md": parse_text,
    ".markdown": parse_text,
    ".txt": parse_text,
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".html": parse_html,
    ".htm": parse_html,
}
