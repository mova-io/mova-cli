#!/usr/bin/env python3
"""Generate realistic multi-format KB fixture files for movate-cli templates.

Produces .txt, .pdf, .docx, .png, .jpg, .tiff files so `mdk kb ingest-all`
exercises every supported parser end-to-end in CI and operator demos.

Run from repo root:
    uv run python scripts/gen_kb_fixtures.py

Requires (all already in dev/optional deps):
    fpdf2, python-docx, Pillow
"""

from __future__ import annotations

import textwrap
from pathlib import Path

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

RAG_QA_KB = Path("src/movate/templates/rag_qa_agent/kb")
HR_POLICY_KB = Path("src/movate/templates/hr_policy_agent/kb")


def _write_txt(path: Path, content: str) -> None:
    path.write_text(textwrap.dedent(content).lstrip(), encoding="utf-8")
    print(f"  wrote {path.relative_to(Path('.'))}")


def _write_pdf(path: Path, title: str, sections: list[tuple[str, str]]) -> None:
    """Generate a multi-page PDF with section headings + body text."""
    import warnings  # noqa: PLC0415

    from fpdf import FPDF  # noqa: PLC0415

    pdf = FPDF()
    pdf.set_margins(left=20, top=20, right=20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    def _latin(s: str) -> str:
        """Replace common Unicode chars that break Helvetica's latin-1 encoding."""
        return (
            s.replace("—", "--")   # em dash
             .replace("–", "-")    # en dash
             .replace("’", "'")    # right single quote
             .replace("‘", "'")    # left single quote
             .replace("“", '"')    # left double quote
             .replace("”", '"')    # right double quote
             .replace("•", "*")    # bullet
             .replace("×", "x")    # multiplication sign
             .replace("é", "e")    # e-acute
        )

    with warnings.catch_warnings():
        warnings.simplefilter("ignore", DeprecationWarning)

        pdf.set_font("Helvetica", style="B", size=16)
        pdf.cell(0, 12, _latin(title), ln=True)
        pdf.ln(4)

        for heading, body in sections:
            pdf.set_font("Helvetica", style="B", size=12)
            pdf.cell(0, 8, _latin(heading), ln=True)
            pdf.set_font("Helvetica", size=10)
            for line in body.strip().splitlines():
                stripped = _latin(line.strip())
                if stripped:
                    pdf.cell(0, 6, stripped, ln=True)
                else:
                    pdf.ln(3)
            pdf.ln(4)

    pdf.output(str(path))
    print(f"  wrote {path.relative_to(Path('.'))}")


def _write_docx(path: Path, title: str, sections: list[tuple[str, str]]) -> None:
    """Generate a DOCX with heading1/body paragraphs."""
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    doc = Document()
    doc.add_heading(title, level=0)

    for heading, body in sections:
        doc.add_heading(heading, level=1)
        for para in body.strip().splitlines():
            p = doc.add_paragraph(para.strip())
            p.style.font.size = Pt(11)

    doc.save(str(path))
    print(f"  wrote {path.relative_to(Path('.'))}")


def _write_image(path: Path, lines: list[str], fmt: str) -> None:
    """Render lines of text onto a white image and save as fmt."""
    from PIL import Image, ImageDraw, ImageFont  # noqa: PLC0415

    width, height = 900, 600
    img = Image.new("RGB", (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)

    # Try to use a legible system font; fall back to default.
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 22)
        font_heading = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 28)
    except OSError:
        try:
            font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 22)
            font_heading = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28
            )
        except OSError:
            font = ImageFont.load_default()
            font_heading = font

    y = 40
    for i, line in enumerate(lines):
        f = font_heading if i == 0 else font
        draw.text((40, y), line, fill=(20, 20, 20), font=f)
        y += 42 if i == 0 else 34

    kwargs: dict = {}
    if fmt.upper() == "TIFF":
        kwargs["compression"] = "tiff_lzw"
    img.save(str(path), format=fmt, **kwargs)
    print(f"  wrote {path.relative_to(Path('.'))}")


# ---------------------------------------------------------------------------
# rag-qa KB files
# ---------------------------------------------------------------------------


def gen_rag_qa_txt() -> None:
    _write_txt(
        RAG_QA_KB / "support-sla.txt",
        """
        Support SLA — Response Time Commitments
        ========================================

        Starter tier
          - First response: within 48 business hours
          - Resolved: within 10 business days
          - Channels: email only (support@example.com)
          - Coverage: Monday–Friday 09:00–17:00 UTC

        Growth tier
          - First response: within 24 business hours
          - Resolved: within 5 business days
          - Channels: email + in-app chat
          - Coverage: Monday–Friday 08:00–20:00 UTC

        Enterprise tier
          - First response: within 4 hours (including weekends)
          - Resolved: within 1 business day for P1, 3 days for P2/P3
          - Channels: email, chat, dedicated Slack channel, phone
          - Coverage: 24 × 7 × 365
          - Named Customer Success Manager assigned on contract signing

        SLA credits
          If response SLAs are missed, customers may request a service credit
          equal to 5% of their monthly invoice for each business day the SLA
          is exceeded, up to a maximum of 30% per invoice period.

        Escalation path
          1. Front-line support agent (first contact)
          2. Senior Support Engineer (if unresolved after 2 business days)
          3. Engineering on-call (P1 incidents, < 4 hour breach)
          4. VP Engineering (P1 incidents breaching 8 hours)

        Priority definitions
          P1 — Service unavailable or data loss in progress
          P2 — Major feature broken, no workaround
          P3 — Minor feature degraded, workaround exists
          P4 — Cosmetic or documentation issue
        """,
    )


def gen_rag_qa_pdf() -> None:
    _write_pdf(
        RAG_QA_KB / "api-rate-limits.pdf",
        title="API Rate Limits & Quotas",
        sections=[
            (
                "Per-Tier Rate Limits",
                """
                Starter tier:    100 requests per minute, 5,000 per day
                Growth tier:     500 requests per minute, 50,000 per day
                Enterprise tier: 2,000 requests per minute, unlimited per day
                All tiers share a burst allowance of 2x the per-minute limit
                for up to 10 seconds before throttling applies.
                """,
            ),
            (
                "Rate Limit Headers",
                """
                Every API response includes:
                  X-RateLimit-Limit:     your per-minute ceiling
                  X-RateLimit-Remaining: requests left this window
                  X-RateLimit-Reset:     Unix timestamp when the window resets
                When you hit the limit the response is HTTP 429 with a
                Retry-After header giving seconds until the window resets.
                """,
            ),
            (
                "Concurrency Limits",
                """
                Starter:    5 concurrent requests
                Growth:     20 concurrent requests
                Enterprise: 100 concurrent requests (configurable)
                Concurrent requests beyond the limit receive HTTP 429
                with a Retry-After of 1 second.
                """,
            ),
            (
                "Token Quotas (LLM endpoints)",
                """
                Starter:    500,000 input tokens / month
                Growth:     5,000,000 input tokens / month
                Enterprise: Negotiated — contact sales@example.com
                Output tokens are not counted against the quota.
                Quota resets on the 1st of each calendar month UTC.
                Overage: requests above quota return HTTP 402 with a
                link to the upgrade page.
                """,
            ),
            (
                "Exemptions",
                """
                Webhook delivery callbacks are not subject to rate limits.
                Internal service-to-service calls using machine API keys
                (role: service_account) are exempt from per-minute limits
                but are capped at the daily maximum.
                """,
            ),
        ],
    )


def gen_rag_qa_docx() -> None:
    _write_docx(
        RAG_QA_KB / "webhook-integration-guide.docx",
        title="Webhook Integration Guide",
        sections=[
            (
                "Overview",
                """
                Webhooks push real-time event notifications to your server
                over HTTPS. Configure them in Settings → Integrations → Webhooks.
                Each event type can be sent to a different URL.
                """,
            ),
            (
                "Registering an Endpoint",
                """
                POST /api/v1/webhooks with:
                  url (string, required): your HTTPS endpoint
                  events (array, required): e.g. ["run.completed", "eval.failed"]
                  secret (string, optional): used to sign payloads
                The endpoint must respond 200 within 5 seconds or the
                delivery is marked failed and retried.
                """,
            ),
            (
                "Signature Verification",
                """
                Every POST includes the header X-Movate-Signature-256 with
                value "sha256=<hex_digest>". Compute HMAC-SHA256 of the raw
                request body using your webhook secret as the key. Compare
                with the header value using a constant-time comparison to
                prevent timing attacks. Reject requests where signatures
                do not match.
                """,
            ),
            (
                "Retry Policy",
                """
                Failed deliveries (non-2xx or timeout) are retried with
                exponential backoff: 30s, 5m, 30m, 2h, 8h (5 attempts total).
                After 5 failures the webhook is disabled and an email is
                sent to the account owner. Re-enable it from Settings.
                """,
            ),
            (
                "Event Payload Shape",
                """
                All events share the envelope:
                  id:         unique event UUID
                  type:       event name (e.g. "run.completed")
                  created_at: ISO-8601 UTC timestamp
                  data:       event-specific payload (see API reference)
                Payloads are delivered as application/json.
                Maximum payload size is 512 KB.
                """,
            ),
            (
                "Supported Event Types",
                """
                run.completed      - agent run finished (success or error)
                run.failed         - agent run failed after all retries
                eval.completed     - eval suite finished
                eval.failed        - eval suite below gate threshold
                kb.ingest.done     - KB ingestion job finished
                billing.quota.80   - token quota at 80% consumed
                billing.quota.100  - token quota exhausted (API will return 402)
                """,
            ),
        ],
    )


def gen_rag_qa_png() -> None:
    _write_image(
        RAG_QA_KB / "uptime-sla.png",
        lines=[
            "Uptime SLA Certificate",
            "99.9% monthly uptime guaranteed (Starter and above)",
            "99.95% monthly uptime guaranteed (Enterprise)",
            "",
            "SLA Credits:",
            "  < 99.9% uptime:  10% of monthly invoice",
            "  < 99.0% uptime:  25% of monthly invoice",
            "  < 95.0% uptime:  50% of monthly invoice",
            "",
            "Uptime is measured at 1-minute intervals from 3 global probes.",
            "Scheduled maintenance windows (announced 72h in advance)",
            "are excluded from uptime calculations.",
            "Status page: status.example.com",
        ],
        fmt="PNG",
    )


def gen_rag_qa_jpg() -> None:
    _write_image(
        RAG_QA_KB / "cancellation-policy.jpg",
        lines=[
            "Cancellation & Refund Policy",
            "Monthly plans: Cancel anytime. Access continues until",
            "  end of the current billing cycle. No partial refunds.",
            "",
            "Annual plans (paid upfront):",
            "  Cancel within 30 days → full refund, no questions asked.",
            "  Cancel after 30 days → prorated refund for unused months,",
            "  less a 10% early-termination fee.",
            "",
            "Enterprise contracts: terms per signed agreement;",
            "  contact your CSM for early-exit options.",
            "",
            "Refunds are issued to the original payment method",
            "within 5-10 business days of approval.",
            "Dispute window: 60 days from invoice date.",
        ],
        fmt="JPEG",
    )


def gen_rag_qa_tiff() -> None:
    _write_image(
        RAG_QA_KB / "compliance-certifications.tiff",
        lines=[
            "Compliance & Security Certifications",
            "SOC 2 Type II  — audited annually by Schellman & Co.",
            "ISO 27001      — certified since 2023, renewal 2026",
            "GDPR           — DPA available; EU data residency option",
            "CCPA           — privacy controls self-certified",
            "HIPAA          — Business Associate Agreement (BAA) available",
            "               on Enterprise tier (request via legal@example.com)",
            "",
            "Penetration testing: annual third-party pentest + quarterly",
            "automated scans. Latest summary report available on request",
            "under NDA for Enterprise customers.",
            "",
            "Data encryption: AES-256 at rest, TLS 1.3 in transit.",
            "Encryption keys managed in AWS KMS with 90-day rotation.",
        ],
        fmt="TIFF",
    )


# ---------------------------------------------------------------------------
# hr-policy KB files
# ---------------------------------------------------------------------------


def gen_hr_txt() -> None:
    _write_txt(
        HR_POLICY_KB / "parking-and-transit-benefits.txt",
        """
        Commuter Benefits — Parking & Transit
        ======================================

        Transit subsidy
          Full-time employees receive a pre-tax transit benefit of up to
          $315 per month (IRS 2026 limit) toward eligible mass transit
          expenses: subway, bus, light rail, commuter rail, vanpool.
          Funds are loaded to your Commuter Benefits Card by the 25th
          of the preceding month. Unused funds roll over for up to
          12 months, then expire.

        Parking subsidy
          Employees who commute by personal vehicle may claim up to
          $315 per month (IRS 2026 limit) in pre-tax parking benefits.
          Eligible: qualified parking at or near the workplace, or at
          a transit facility for a park-and-ride commute.

        How to enroll
          1. Log in to BambooHR → Benefits → Commuter Benefits.
          2. Select Transit, Parking, or both.
          3. Enter your monthly election amount (up to the IRS limit).
          4. Changes take effect the 1st of the following month.
          5. Enrollment open year-round (no open-enrollment window).

        Remote employees
          Employees who are 100% remote and have no designated office
          are not eligible for the commuter benefit. Employees on a
          hybrid schedule (any in-office days) are fully eligible.

        Eligible providers (transit card)
          Clipper (Bay Area), Metro SmarTrip (DC), CharlieCard (Boston),
          ORCA (Seattle), TAP (LA), Ventra (Chicago), OMNY (NYC), and
          all NTD-registered transit agencies.
        """,
    )


def gen_hr_pdf() -> None:
    _write_pdf(
        HR_POLICY_KB / "expense-reimbursement-policy.pdf",
        title="Employee Expense Reimbursement Policy",
        sections=[
            (
                "Eligible Expense Categories",
                """
                Business travel: flights (economy), hotels ($250/night cap in
                major metros, $175 elsewhere), ground transport, meals ($75/day
                per diem while traveling — no receipts required under $25).
                Client entertainment: pre-approved by manager, max $150/person.
                Home office: one-time $500 setup allowance for new hires,
                $100/year ongoing (keyboard, mouse, monitor accessories).
                Professional development: books, online courses — up to
                $1,500/year without additional approval; above that needs VP sign-off.
                """,
            ),
            (
                "Receipt Requirements",
                """
                Receipts are required for all individual expenses of $25 or more.
                Receipts must show: vendor name, date, itemized amount, and
                payment method (last 4 digits). Screenshots of digital receipts
                are acceptable. Credit card statements alone are not sufficient.
                Missing receipts: submit a Missing Receipt Affidavit (HR portal)
                for expenses up to $100. Above $100 requires manager approval.
                """,
            ),
            (
                "Submission Deadlines",
                """
                Submit expenses within 60 days of the expense date.
                Expenses older than 60 days require VP Finance approval.
                Expenses older than 90 days will not be reimbursed.
                Year-end cutoff: all expenses for a calendar year must be
                submitted by January 15 of the following year.
                """,
            ),
            (
                "Manager Approval Thresholds",
                """
                Up to $500: direct manager approval only.
                $501–$2,000: direct manager + skip-level approval.
                $2,001–$5,000: VP-level approval required.
                Above $5,000: CFO approval required.
                Client entertainment always requires manager pre-approval
                before the expense is incurred, regardless of amount.
                """,
            ),
            (
                "Reimbursement Processing",
                """
                Approved expenses are processed in the bi-weekly payroll
                run (same cycle as salary). Submit by Thursday midnight for
                the following Friday's payroll. Expenses submitted after
                the cutoff roll to the next bi-weekly cycle.
                Reimbursements appear as a separate line item on your pay stub
                labeled EXPENSE REIMB.
                """,
            ),
        ],
    )


def gen_hr_docx() -> None:
    _write_docx(
        HR_POLICY_KB / "performance-review-calendar.docx",
        title="Performance Review Calendar & Process",
        sections=[
            (
                "Annual Review Cycle Overview",
                """
                The company runs a single annual performance review cycle
                aligned to the calendar year (Jan 1 – Dec 31).
                Reviews have four phases: self-assessment, manager review,
                calibration, and compensation planning.
                """,
            ),
            (
                "Phase 1: Self-Assessment (Q1 — January 15–31)",
                """
                All employees complete a self-assessment in Lattice by January 31.
                Cover: top 3 accomplishments, areas for growth, goal progress,
                and a proposed rating (Exceeds / Meets / Developing / Underperforming).
                Self-assessments are visible to your manager only, not peers.
                """,
            ),
            (
                "Phase 2: Manager Review (Q1 — February 1–28)",
                """
                Managers write reviews for each direct report using the
                self-assessment as input. Managers propose ratings and draft
                talking points for the 1:1 review conversation.
                360 peer feedback (optional, 2–3 peers selected by employee)
                is collected Feb 1–14 and shared with managers Feb 15.
                """,
            ),
            (
                "Phase 3: Calibration (Q2 — March)",
                """
                Directors and VPs hold calibration sessions to normalize ratings
                across teams. The calibration guideline targets:
                  Exceeds:           ~15% of employees
                  Meets:             ~70% of employees
                  Developing:        ~12% of employees
                  Underperforming:   ~3% of employees
                Calibration outcomes are final; managers communicate results
                in 1:1s by March 31.
                """,
            ),
            (
                "Phase 4: Compensation Planning (Q2 — April)",
                """
                Comp adjustments (merit increases, equity refreshes) are
                finalized in April based on calibrated ratings:
                  Exceeds:       4–6% merit increase + equity refresh
                  Meets:         2–4% merit increase
                  Developing:    0–2% merit increase, no equity refresh
                  Underperforming: 0%; Performance Improvement Plan initiated
                New compensation effective May 1 payroll.
                """,
            ),
            (
                "Mid-Year Check-In (Q3 — July)",
                """
                A lightweight mid-year check-in is conducted in July.
                This is NOT a formal rating exercise. Managers update goals
                in Lattice and have a 30-minute 1:1 focused on H2 priorities.
                No calibration or comp changes occur at mid-year.
                """,
            ),
        ],
    )


def gen_hr_png() -> None:
    _write_image(
        HR_POLICY_KB / "employee-assistance-program.png",
        lines=[
            "Employee Assistance Program (EAP)",
            "Provider: Lyra Health  |  Available: 24 hours / 7 days / 365 days",
            "",
            "What is included (all FREE and CONFIDENTIAL):",
            "  8 counseling sessions per year (in-person or video)",
            "  Unlimited mental health coaching sessions",
            "  Financial counseling: 3 sessions per topic per year",
            "  Legal consultation: 30 min per issue, no charge",
            "  Work-life balance support: childcare, eldercare referrals",
            "",
            "How to access:",
            "  Web:   app.lyrahealth.com/example-company",
            "  Phone: 1-877-EAP-LYRA (1-877-327-5972)",
            "  App:   Lyra Health iOS / Android",
            "",
            "All sessions are 100% confidential. Usage is never reported",
            "to HR or your manager. Covered for employee + household members.",
        ],
        fmt="PNG",
    )


def gen_hr_jpg() -> None:
    _write_image(
        HR_POLICY_KB / "company-holidays-2026.jpg",
        lines=[
            "Company Holidays 2026  (all employees, all offices)",
            "Jan 1   New Year's Day",
            "Jan 19  Martin Luther King Jr. Day",
            "Feb 16  Presidents' Day",
            "May 25  Memorial Day",
            "Jun 19  Juneteenth National Independence Day",
            "Jul 4   Independence Day (US offices) / Canada Day Jul 1",
            "Sep 7   Labor Day",
            "Nov 26  Thanksgiving Day",
            "Nov 27  Day after Thanksgiving (floating)",
            "Dec 24  Christmas Eve (half day from noon local time)",
            "Dec 25  Christmas Day",
            "Dec 31  New Year's Eve (half day from noon local time)",
            "",
            "Total: 13 company holidays in 2026.",
            "Non-US offices: local statutory holidays substituted where applicable.",
        ],
        fmt="JPEG",
    )


def gen_hr_tiff() -> None:
    _write_image(
        HR_POLICY_KB / "pay-stub-guide.tiff",
        lines=[
            "Understanding Your Pay Stub",
            "Gross Pay     Total earnings before any deductions.",
            "              Includes base salary, OT, bonuses, commissions.",
            "Federal Tax   Federal income tax withheld (W-4 elections apply).",
            "State Tax     State income tax withheld (varies by work location).",
            "OASDI         Social Security: 6.2% of gross up to $176,100 (2026).",
            "Medicare      Hospital Insurance: 1.45% of gross (no wage base cap).",
            "401(k)        Your elected pre-tax or Roth deferral amount.",
            "Med/Den/Vis   Health, dental, vision premium (pre-tax).",
            "FSA / HSA     Healthcare or dependent-care FSA election.",
            "Net Pay       Take-home after ALL deductions. Wire to your bank.",
            "YTD           Year-to-date totals for each line item — resets Jan 1.",
            "",
            "Questions: payroll@example.com or HR portal → Payroll → Inquiries",
        ],
        fmt="TIFF",
    )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating rag-qa KB fixtures...")
    gen_rag_qa_txt()
    gen_rag_qa_pdf()
    gen_rag_qa_docx()
    gen_rag_qa_png()
    gen_rag_qa_jpg()
    gen_rag_qa_tiff()

    print("\nGenerating hr-policy KB fixtures...")
    gen_hr_txt()
    gen_hr_pdf()
    gen_hr_docx()
    gen_hr_png()
    gen_hr_jpg()
    gen_hr_tiff()

    print("\nDone. Run `mdk kb ingest-all --dry-run` to verify discovery.")
