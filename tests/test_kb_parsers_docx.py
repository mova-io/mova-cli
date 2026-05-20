"""Tests for ``movate.kb.parsers.parse_docx`` (PR-L).

Same shape as the PDF tests in ``test_kb_parsers.py``:

* Extension dispatch routes .docx to parse_docx
* Round-trip: docx with N paragraphs → N chunks joined by \\n\\n
* Headings get preserved as their own paragraph break
* Defensive returns: corrupt bytes, empty document → None

Synthetic .docx files built inline via python-docx so we don't ship
a binary fixture and the test stays hermetic.
"""

from __future__ import annotations

import io

import docx
import pytest

from movate.kb.parsers import (
    SUPPORTED_EXTENSIONS,
    is_supported_extension,
    parse_document,
    parse_docx,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_docx_bytes(*paragraphs: tuple[str, str] | str) -> bytes:
    """Build a synthetic .docx in memory.

    Each item is either:
    * A plain string → added as a Normal-style paragraph.
    * A ``(style, text)`` tuple → ``style`` is the Word style name
      (e.g. ``"Heading 1"``, ``"Heading 2"``) and ``text`` is the
      paragraph body.
    """
    buf = io.BytesIO()
    doc = docx.Document()
    for item in paragraphs:
        if isinstance(item, tuple):
            style, text = item
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(item)
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Extension dispatch
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_docx_in_supported_extensions() -> None:
    """``.docx`` is in the public supported set after PR-L."""
    assert ".docx" in SUPPORTED_EXTENSIONS
    assert is_supported_extension("policy.DOCX")  # case-insensitive
    # Legacy binary .doc explicitly NOT supported.
    assert not is_supported_extension("legacy.doc")


@pytest.mark.unit
def test_parse_document_routes_docx_to_docx_parser() -> None:
    """parse_document(filename.docx) → parse_docx. Single integration
    point for endpoint + CLI."""
    body = _make_docx_bytes("Hello DOCX world")
    out = parse_document("policy.docx", body)
    assert out is not None
    assert "Hello DOCX world" in out


# ---------------------------------------------------------------------------
# parse_docx — happy paths
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_parse_docx_extracts_single_paragraph() -> None:
    body = _make_docx_bytes("This is a refund policy paragraph.")
    out = parse_docx(body)
    assert out is not None
    assert "refund policy" in out


@pytest.mark.unit
def test_parse_docx_joins_paragraphs_with_paragraph_breaks() -> None:
    """Multi-paragraph round-trip preserves the paragraph boundary
    via ``\\n\\n`` so the chunker treats them as natural breaks."""
    body = _make_docx_bytes(
        "First paragraph about refunds.",
        "Second paragraph about cancellations.",
    )
    out = parse_docx(body)
    assert out is not None
    assert "First paragraph" in out
    assert "Second paragraph" in out
    assert "\n\n" in out  # paragraph boundary preserved


@pytest.mark.unit
def test_parse_docx_includes_heading_text() -> None:
    """Heading-style paragraphs survive the round-trip with their
    text intact — operators rely on heading content as semantic
    anchors during retrieval."""
    body = _make_docx_bytes(
        ("Heading 1", "Refund Policy"),
        "Annual subscriptions refundable within 14 days.",
    )
    out = parse_docx(body)
    assert out is not None
    assert "Refund Policy" in out
    assert "Annual subscriptions" in out


@pytest.mark.unit
def test_parse_docx_skips_empty_paragraphs() -> None:
    """Empty paragraphs (blank lines in the source doc) don't
    contribute spurious empty chunks."""
    body = _make_docx_bytes(
        "Paragraph one.",
        "",  # empty
        "   ",  # whitespace-only
        "Paragraph two.",
    )
    out = parse_docx(body)
    assert out is not None
    # Two paragraphs joined by a single \n\n — not three with stray
    # blank chunks in between.
    assert out.count("\n\n") == 1


# ---------------------------------------------------------------------------
# Defensive returns
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_parse_docx_returns_none_on_corrupt_bytes() -> None:
    """Random bytes / not-a-docx → None. Operator sees ``status="skipped"``
    in the per-file upload result, batch isn't sunk."""
    assert parse_docx(b"not a docx file at all") is None
    assert parse_docx(b"") is None


@pytest.mark.unit
def test_parse_docx_returns_none_on_empty_document() -> None:
    """A .docx with no paragraphs (or all empty) → None.
    Embedded-images-only docs hit this path."""
    buf = io.BytesIO()
    doc = docx.Document()
    # Save with no paragraphs added.
    doc.save(buf)
    out = parse_docx(buf.getvalue())
    assert out is None


@pytest.mark.unit
def test_parse_document_docx_failure_returns_none() -> None:
    """End-to-end: .docx with junk content returns None at dispatch
    level. Callers report ``status="skipped"`` cleanly."""
    assert parse_document("policy.docx", b"corrupt") is None


# ---------------------------------------------------------------------------
# find_files dispatch — picks up .docx files
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_find_files_includes_docx(tmp_path: object) -> None:
    """``mdk kb ingest <dir>`` walks .docx alongside .md/.txt/.pdf.
    Regression guard for the SUPPORTED_EXTENSIONS extension."""
    from pathlib import Path  # noqa: PLC0415

    from movate.kb.ingest import find_files  # noqa: PLC0415

    root = Path(str(tmp_path))
    (root / "a.md").write_text("md content")
    (root / "b.docx").write_bytes(_make_docx_bytes("docx content"))
    (root / "skip.html").write_text("html not supported")

    found = {p.name for p in find_files(root)}
    assert "a.md" in found
    assert "b.docx" in found
    assert "skip.html" not in found
